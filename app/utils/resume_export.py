"""Utilities for exporting resumes to PDF and DOCX formats"""

import io
from typing import BinaryIO
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_resume_pdf(resume_text: str, name: str) -> bytes:
    """
    Generate a PDF file from resume text
    
    Args:
        resume_text: The formatted resume text
        name: Candidate name for the document
        
    Returns:
        bytes: PDF file content
    """
    buffer = io.BytesIO()
    
    # Create PDF document
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.75*inch,
        leftMargin=0.75*inch,
        topMargin=0.75*inch,
        bottomMargin=0.75*inch
    )
    
    # Container for the 'Flowable' objects
    elements = []
    
    # Define styles
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=RGBColor(0, 0, 0),
        spaceAfter=12,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=RGBColor(0, 0, 0),
        spaceAfter=6,
        spaceBefore=12,
        fontName='Helvetica-Bold'
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['BodyText'],
        fontSize=11,
        textColor=RGBColor(0, 0, 0),
        spaceAfter=6,
        alignment=TA_LEFT,
        fontName='Helvetica'
    )
    
    # Parse resume text and add to PDF
    lines = resume_text.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            elements.append(Spacer(1, 0.1*inch))
            continue
            
        # Check if line is a title (all caps or starts with specific patterns)
        if line.isupper() and len(line) > 3:
            # Section heading
            elements.append(Paragraph(line, heading_style))
        elif line.startswith('=') or line.startswith('-'):
            # Separator line - skip
            continue
        else:
            # Regular text - escape special characters for reportlab
            safe_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            elements.append(Paragraph(safe_line, body_style))
    
    # Build PDF
    doc.build(elements)
    
    # Get the value of the BytesIO buffer
    pdf_content = buffer.getvalue()
    buffer.close()
    
    return pdf_content


def generate_resume_docx(resume_text: str, name: str) -> bytes:
    """
    Generate a DOCX file from resume text
    
    Args:
        resume_text: The formatted resume text
        name: Candidate name for the document
        
    Returns:
        bytes: DOCX file content
    """
    # Create a new Document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Parse resume text and add to document
    lines = resume_text.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            doc.add_paragraph()
            continue
        
        # Check if line is a section heading (all caps)
        if line.isupper() and len(line) > 3 and not line.startswith('=') and not line.startswith('-'):
            # Add section heading
            heading = doc.add_heading(line, level=2)
            heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        elif line.startswith('=') or line.startswith('-'):
            # Separator line - skip or add horizontal line
            continue
        else:
            # Regular paragraph
            p = doc.add_paragraph(line)
            p.style = 'Normal'
    
    # Save to BytesIO buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    docx_content = buffer.read()
    buffer.close()
    
    return docx_content
